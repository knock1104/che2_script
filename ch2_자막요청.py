import streamlit as st
from datetime import date
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor

st.set_page_config(page_title="자막 업로더", page_icon="📝", layout="centered")

# ----------------------
# Utilities
# ----------------------
PARTS = [
    "소프라노", "알토", "테너", "베이스", "여성파트", "남성파트", "투티", "합창", "전주", "간주"
]

PART_COLORS = {
    "소프라노": (220, 20, 60),
    "알토": (30, 144, 255),
    "테너": (46, 139, 87),
    "베이스": (147, 112, 219),
    "여성파트": (255, 0, 255),
    "남성파트": (0, 128, 128),
    "투티": (0, 0, 0),            # black
    "합창": (0, 0, 0),            # black
    "전주": (112, 128, 144),
    "간주": (255, 140, 0),
}

@st.cache_data(show_spinner=False)
def _get_today():
    return date.today()

# ----------------------
# Session State Defaults
# ----------------------
if "authed" not in st.session_state:
    st.session_state.authed = False
if "author" not in st.session_state:
    st.session_state.author = ""
if "entries" not in st.session_state:
    # Each entry: {"text": str, "parts": list[str]}
    st.session_state.entries = []
if "meta" not in st.session_state:
    st.session_state.meta = {
        "date": _get_today(),
        "singer": "",
        "has_part": "구분 없음",
    }

# ----------------------
# Helper Actions
# ----------------------

def add_entry(text: str, parts: list[str]):
    st.session_state.entries.append({"text": text.strip(), "parts": parts[:]})

def move_up(idx: int):
    if idx > 0:
        items = st.session_state.entries
        items[idx-1], items[idx] = items[idx], items[idx-1]

def move_down(idx: int):
    items = st.session_state.entries
    if idx < len(items) - 1:
        items[idx+1], items[idx] = items[idx], items[idx+1]

def delete_row(idx: int):
    items = st.session_state.entries
    if 0 <= idx < len(items):
        items.pop(idx)

def clear_all():
    st.session_state.entries = []

# ----------------------
# Word Export
# ----------------------

def make_docx(author: str, worship_date: date, singer: str, has_part: bool, entries: list[dict]) -> BytesIO:
    doc = Document()

    # Title & meta
    title = doc.add_paragraph()
    run = title.add_run("찬양 자막")
    run.font.size = Pt(20)
    run.bold = True

    meta = doc.add_paragraph()
    meta.add_run(f"작성자: {author}\t").bold = True
    meta.add_run(f"찬양일: {worship_date.strftime('%Y-%m-%d')}\t").bold = True
    meta.add_run(f"찬양대/특송자: {singer}").bold = True

    doc.add_paragraph("")

    # Legend (when part on)
    if has_part:
        legend = doc.add_paragraph()
        legend.add_run("[파트 색상 안내] ")
        for p in PARTS:
            r, g, b = PART_COLORS[p]
            run = legend.add_run(f"{p} ")
            run.font.color.rgb = RGBColor(r, g, b)
        doc.add_paragraph("")

    # Body lines
    for i, item in enumerate(entries, start=1):
        text = item.get("text", "").strip()
        parts = item.get("parts", []) if has_part else []
        para = doc.add_paragraph()
        prefix = f"{i}. "
        pre = para.add_run(prefix)
        pre.font.size = Pt(12)
        pre.bold = True

        # Color by first part if exists
        run = para.add_run(text)
        run.font.size = Pt(12)
        if parts:
            first = parts[0]
            if first in PART_COLORS:
                r, g, b = PART_COLORS[first]
                run.font.color.rgb = RGBColor(r, g, b)

        if has_part and parts:
            tag = para.add_run("    [" + ", ".join(parts) + "]")
            tag.italic = True

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ----------------------
# Login Page (Landing)
# ----------------------

def login_page():
    st.markdown("## 🔐 로그인")
    st.caption("테스트 페이지 비밀번호는 모든 사용자 공통으로 0001 입니다.")

    author = st.text_input("작성자 이름", key="login_author", placeholder="예) 홍길동")
    pw = st.text_input("비밀번호", type="password", key="login_pw")

    col1, col2 = st.columns(2)
    with col1:
        if st.button("로그인", type="primary"):
            if pw == "0001" and author.strip():
                st.session_state.authed = True
                st.session_state.author = author.strip()
                st.success("로그인 되었습니다.")
                st.rerun()
            else:
                st.error("작성자와 비밀번호(0001)를 확인하세요.")
    with col2:
        st.button("초기화", on_click=lambda: st.session_state.update({"login_author": "", "login_pw": ""}))

# ----------------------
# Main Page
# ----------------------

def main_page():
    st.markdown("## 📝 자막 입력")
    with st.expander("기본 정보", expanded=True):
        st.session_state.meta["date"] = st.date_input("찬양일", value=st.session_state.meta["date"])  # 달력 형식
        st.session_state.meta["singer"] = st.text_input("찬양대 / 특송자", value=st.session_state.meta["singer"], placeholder="예) 할렐루야 찬양대 / 김OO 집사 특송")
        st.session_state.meta["has_part"] = st.selectbox("자막 파트 구분", ["구분 없음", "구분 있음"], index=0 if st.session_state.meta["has_part"]=="구분 없음" else 1)
        st.caption("※ '구분 있음' 선택 시 각 자막 줄에 파트를 지정(복수 선택 가능)하고, 결과 Word에서 파트별 색상으로 표시됩니다.")

    has_part = st.session_state.meta["has_part"] == "구분 있음"

    st.divider()

    # Add entry form
    st.subheader("자막 추가")
    with st.form("add_form", clear_on_submit=True):
        parts = []
        text = ""
        measures = 0

        if has_part:
            parts = st.multiselect("파트 선택(복수 선택 가능)", PARTS, default=[])
            # 전주/간주 선택 시 마디 수 옵션
            if any(p in parts for p in ["전주", "간주"]):
                measures = st.number_input("마디 수(선택)", min_value=0, max_value=128, value=0, step=1, help="0이면 '전주', '간주'만 표기")

            # 전주/간주만(단일 선택)이면서 내용 미입력 시 자동 텍스트
            if parts and set(parts).issubset({"전주", "간주"}):
                # 우선 순위: 전주 -> 간주 (단일선택 기준)
                picked = None
                if "전주" in parts and len(parts) == 1:
                    picked = "전주"
                elif "간주" in parts and len(parts) == 1:
                    picked = "간주"
                if picked:
                    text = f"{picked} {measures}마디" if measures and measures > 0 else picked
                else:
                    # 전주와 간주를 동시에 선택한 경우에는 입력창 열어둠
                    text = st.text_area("자막 내용", placeholder="자막 한 줄을 입력하고 '추가'를 누르세요", height=80)
            else:
                text = st.text_area("자막 내용", placeholder="자막 한 줄을 입력하고 '추가'를 누르세요", height=80)
        else:
            text = st.text_area("자막 내용", placeholder="자막 한 줄을 입력하고 '추가'를 누르세요", height=80)

        submitted = st.form_submit_button("추가", use_container_width=True)
        if submitted:
            # 텍스트 자동 생성 규칙(전주/간주 단일 선택 & 공란인 경우 한번 더 보장)
            if has_part and not text.strip() and parts and set(parts).issubset({"전주", "간주"}) and len(parts) == 1:
                picked = parts[0]
                text = f"{picked} {measures}마디" if measures and measures > 0 else picked

            if not text.strip():
                st.warning("자막 내용을 입력하세요.")
            else:
                add_entry(text, parts if has_part else [])
                st.success("추가되었습니다.")

    # List & reorder
    st.subheader("자막 목록 (순서 조정 가능)")
    if not st.session_state.entries:
        st.info("아직 추가된 자막이 없습니다.")
    else:
        for idx, item in enumerate(st.session_state.entries):
            c1, c2, c3, c4, c5 = st.columns([1, 3, 6, 1, 1])
            with c1:
                st.write(f"**{idx+1}**")
            with c2:
                if has_part:
                    parts_label = ", ".join(item.get("parts", [])) or "(미지정)"
                    st.write(parts_label)
                else:
                    st.write("-")
            with c3:
                st.text_area(label="", value=item.get("text",""), key=f"row_text_{idx}", height=70)
            with c4:
                st.button("▲", key=f"up_{idx}", on_click=move_up, args=(idx,))
                st.button("▼", key=f"down_{idx}", on_click=move_down, args=(idx,))
            with c5:
                st.button("삭제", key=f"del_{idx}", on_click=delete_row, args=(idx,))

        # Apply edited texts back
        for idx in range(len(st.session_state.entries)):
            new_text = st.session_state.get(f"row_text_{idx}", st.session_state.entries[idx]["text"]) or ""
            st.session_state.entries[idx]["text"] = new_text

        colA, colB = st.columns([1, 1])
        with colA:
            st.button("전체 초기화", on_click=clear_all, type="secondary")
        with colB:
            author = st.session_state.author or ""
            buf = make_docx(
                author=author,
                worship_date=st.session_state.meta["date"],
                singer=st.session_state.meta["singer"],
                has_part=has_part,
                entries=st.session_state.entries,
            )
            st.download_button(
                label="Word 파일로 내보내기 (.docx)",
                data=buf,
                file_name=f"찬양자막_{st.session_state.meta['date'].strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )

    st.divider()
    st.caption("Tip: 파트는 복수 선택 가능하며, '전주' 또는 '간주'만 단독 선택하면 입력 없이 자동으로 텍스트가 채워집니다. 마디 수를 입력하면 예: 전주 4마디.")

# ----------------------
# Router
# ----------------------

st.markdown("### che2 자막 업로더")

if not st.session_state.authed:
    login_page()
else:
    st.info(f"접속: {st.session_state.author}")
    if st.button("로그아웃"):
        st.session_state.authed = False
        st.session_state.author = ""
        st.rerun()
    main_page()
